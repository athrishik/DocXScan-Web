#!/usr/bin/env python
# coding: utf-8

"""
DocXScan v3.0 - Professional Document Scanner (Streamlit Edition)
Copyright 2025 Hrishik Kunduru. All rights reserved.

Professional document scanner with intelligent token detection and modern UI.
"""

import streamlit as st
import tempfile
import os
import zipfile
import pandas as pd
from datetime import datetime
from docx import Document
import json
import shutil
from pathlib import Path
import threading
from io import BytesIO
import base64
import time
import glob
import platform

# Configure Streamlit page
st.set_page_config(
    page_title="DocXScan v3.0 Web",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Modern CSS styling
def load_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600;700&display=swap');
    
    :root {
        --bg-primary: #0A0E1A;
        --bg-secondary: #151922;
        --bg-card: #1E2532;
        --accent-blue: #2563EB;
        --accent-green: #10B981;
        --accent-orange: #F59E0B;
        --accent-red: #EF4444;
        --accent-purple: #8B5CF6;
        --text-primary: #FFFFFF;
        --text-secondary: #E5E7EB;
        --text-muted: #9CA3AF;
        --border-color: #374151;
        --success: #10B981;
        --warning: #F59E0B;
        --error: #EF4444;
        --shadow-sm: 0 2px 4px rgba(0, 0, 0, 0.1);
        --shadow-md: 0 4px 12px rgba(0, 0, 0, 0.2);
        --shadow-lg: 0 8px 24px rgba(0, 0, 0, 0.3);
        --shadow-xl: 0 12px 32px rgba(0, 0, 0, 0.4);
    }
    
    .stApp {
        background: linear-gradient(135deg, var(--bg-primary) 0%, var(--bg-secondary) 100%);
        color: var(--text-primary);
        font-family: 'Inter', sans-serif;
    }
    
    /* ============= HEADER STYLING ============= */
    .main-header {
        background: linear-gradient(145deg, var(--bg-card), #2A3441);
        border: 2px solid var(--accent-blue);
        border-radius: 16px;
        padding: 2rem;
        margin-bottom: 2rem;
        text-align: center;
        box-shadow: var(--shadow-lg);
        position: relative;
        overflow: hidden;
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(45deg, transparent 30%, rgba(37, 99, 235, 0.05) 50%, transparent 70%);
        animation: shimmer 3s infinite;
    }
    
    @keyframes shimmer {
        0% { transform: translateX(-100%); }
        100% { transform: translateX(100%); }
    }
    
    .main-title {
        font-size: 2.5rem;
        font-weight: 800;
        color: var(--accent-blue);
        margin-bottom: 0.5rem;
        text-shadow: 0 0 20px rgba(37, 99, 235, 0.3);
        position: relative;
        z-index: 1;
    }
    
    .main-subtitle {
        font-size: 1.1rem;
        color: var(--text-secondary);
        font-weight: 500;
        margin: 0;
        position: relative;
        z-index: 1;
    }
    
    /* ============= CARD SYSTEM ============= */
    .modern-card {
        background: linear-gradient(145deg, var(--bg-card), #242B3D);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        box-shadow: var(--shadow-md);
        transition: all 0.3s ease;
    }
    
    .modern-card:hover {
        border-color: var(--accent-blue);
        box-shadow: var(--shadow-lg);
    }
    
    .card-title {
        font-size: 1.2rem;
        font-weight: 600;
        color: var(--text-primary);
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    .status-indicator {
        width: 8px;
        height: 8px;
        background: var(--success);
        border-radius: 50%;
        margin-left: auto;
        animation: pulse 2s infinite;
    }
    
    @keyframes pulse {
        0% { opacity: 1; }
        50% { opacity: 0.5; }
        100% { opacity: 1; }
    }
    
    /* ============= ENHANCED FOLDER BROWSER ============= */
    .folder-browser-enhanced {
        background: linear-gradient(145deg, #1A1A1A, #222222);
        border: 3px dashed var(--accent-blue);
        border-radius: 16px;
        padding: 30px;
        text-align: center;
        transition: all 0.3s ease;
        margin: 15px 0;
        cursor: pointer;
        position: relative;
        overflow: hidden;
    }
    
    .folder-browser-enhanced::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(45deg, transparent 30%, rgba(37, 99, 235, 0.1) 50%, transparent 70%);
        animation: shimmer 2s infinite;
    }
    
    .folder-browser-enhanced:hover {
        border-color: var(--accent-green);
        background: linear-gradient(145deg, #222222, #2A2A2A);
        transform: translateY(-4px);
        box-shadow: var(--shadow-lg);
    }
    
    .folder-browser-icon {
        font-size: 3rem;
        color: var(--accent-blue);
        margin-bottom: 15px;
        animation: bounce 2s infinite;
        position: relative;
        z-index: 1;
    }
    
    @keyframes bounce {
        0%, 20%, 50%, 80%, 100% { transform: translateY(0); }
        40% { transform: translateY(-10px); }
        60% { transform: translateY(-5px); }
    }
    
    .folder-browser-title {
        color: var(--text-primary);
        font-weight: 700;
        font-size: 1.2rem;
        margin-bottom: 8px;
        position: relative;
        z-index: 1;
    }
    
    .folder-browser-subtitle {
        color: var(--text-secondary);
        font-weight: 500;
        font-size: 0.9rem;
        position: relative;
        z-index: 1;
    }
    
    .selected-folder-display {
        background: linear-gradient(145deg, var(--success), #059669);
        border: 2px solid var(--success);
        border-radius: 12px;
        padding: 20px;
        margin: 15px 0;
        text-align: center;
        box-shadow: var(--shadow-md);
    }
    
    .selected-folder-text {
        color: white;
        font-weight: 600;
        font-size: 1rem;
        margin-bottom: 5px;
    }
    
    .selected-folder-path {
        color: rgba(255, 255, 255, 0.9);
        font-size: 0.85rem;
        word-break: break-all;
        font-family: 'JetBrains Mono', monospace;
    }
    
    /* ============= TAB SYSTEM ENHANCEMENT ============= */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: transparent !important;
        border-bottom: none !important;
        margin-bottom: 1rem;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: linear-gradient(145deg, var(--bg-card), #2A3441) !important;
        color: var(--text-secondary) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        padding: 10px 20px !important;
        margin: 0 !important;
        transition: all 0.3s ease !important;
        box-shadow: var(--shadow-sm) !important;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background: linear-gradient(145deg, #2A3441, var(--bg-card)) !important;
        border-color: var(--accent-blue) !important;
        transform: translateY(-1px) !important;
        box-shadow: var(--shadow-md) !important;
    }
    
    .stTabs [aria-selected="true"][data-baseweb="tab"] {
        background: linear-gradient(135deg, var(--accent-blue), var(--accent-green)) !important;
        color: white !important;
        border-color: var(--accent-blue) !important;
        box-shadow: var(--shadow-md) !important;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.3) !important;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        padding: 0 !important;
    }
    
    /* ============= CONSOLE STYLING ============= */
    .console-area {
        background: #000000;
        color: #00FF88;
        font-family: 'JetBrains Mono', monospace;
        border: 2px solid #00FF88;
        border-radius: 8px;
        padding: 1rem;
        height: 300px;
        overflow-y: auto;
        font-size: 0.9rem;
        line-height: 1.4;
        box-shadow: var(--shadow-md);
        position: relative;
    }
    
    .console-area::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(45deg, transparent 30%, rgba(0, 255, 136, 0.05) 50%, transparent 70%);
        animation: scan-line 2s infinite;
        pointer-events: none;
    }
    
    @keyframes scan-line {
        0% { transform: translateY(-100%); }
        100% { transform: translateY(100%); }
    }
    
    /* ============= METRICS AND CARDS ============= */
    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin: 1rem 0;
    }
    
    .metric-card {
        background: linear-gradient(145deg, var(--bg-card), #242B3D);
        border: 1px solid var(--border-color);
        border-radius: 8px;
        padding: 1rem;
        text-align: center;
        transition: all 0.3s ease;
        box-shadow: var(--shadow-sm);
    }
    
    .metric-card:hover {
        transform: translateY(-4px);
        border-color: var(--accent-blue);
        box-shadow: var(--shadow-md);
    }
    
    .metric-value {
        font-size: 1.5rem;
        font-weight: 700;
        color: var(--accent-blue);
        margin-bottom: 0.25rem;
        text-shadow: 0 0 10px rgba(37, 99, 235, 0.3);
    }
    
    .metric-label {
        font-size: 0.875rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 0.5px;
        font-weight: 500;
    }
    
    /* ============= SIDEBAR STYLING ============= */
    .css-1d391kg, section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, var(--bg-secondary) 0%, #1A2332 100%) !important;
        border-right: 2px solid var(--accent-blue) !important;
        box-shadow: var(--shadow-lg) !important;
    }
    
    /* Sidebar Text - High Visibility */
    .css-1d391kg *, section[data-testid="stSidebar"] *,
    section[data-testid="stSidebar"] .stMarkdown *,
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] .stSelectbox label,
    section[data-testid="stSidebar"] .stTextInput label,
    section[data-testid="stSidebar"] .stTextArea label,
    section[data-testid="stSidebar"] .stFileUploader label {
        color: var(--text-primary) !important;
        font-weight: 500 !important;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.8) !important;
    }
    
    /* Sidebar Headers */
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] h4,
    .css-1d391kg h1, .css-1d391kg h2, .css-1d391kg h3, .css-1d391kg h4 {
        color: var(--accent-blue) !important;
        font-weight: 700 !important;
        text-shadow: 0 0 10px rgba(37, 99, 235, 0.5) !important;
        margin-bottom: 1rem !important;
    }
    
    /* ============= INPUT FIELD ENHANCEMENTS ============= */
    /* Text Input Fields */
    .stTextInput > div > div > input,
    section[data-testid="stSidebar"] .stTextInput > div > div > input {
        background-color: #1A1A1A !important;
        color: var(--text-primary) !important;
        border: 2px solid var(--accent-blue) !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 12px !important;
        outline: none !important;
        box-shadow: var(--shadow-sm) !important;
        transition: all 0.3s ease !important;
        font-family: 'JetBrains Mono', monospace !important;
    }
    
    .stTextInput > div > div > input::placeholder {
        color: var(--text-muted) !important;
        opacity: 0.8 !important;
        font-style: italic !important;
    }
    
    .stTextInput > div > div > input:focus {
        outline: none !important;
        box-shadow: 0 0 0 3px rgba(16, 185, 129, 0.2) !important;
        border: 2px solid var(--accent-green) !important;
        transform: translateY(-1px) !important;
    }
    
    /* Text Area Fields */
    .stTextArea > div > div > textarea,
    section[data-testid="stSidebar"] .stTextArea > div > div > textarea {
        background-color: #1A1A1A !important;
        color: var(--text-primary) !important;
        border: 2px solid var(--accent-blue) !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 12px !important;
        outline: none !important;
        box-shadow: var(--shadow-sm) !important;
        font-family: 'JetBrains Mono', monospace !important;
        transition: all 0.3s ease !important;
    }
    
    .stTextArea > div > div > textarea::placeholder {
        color: var(--text-muted) !important;
        opacity: 0.8 !important;
        font-style: italic !important;
    }
    
    .stTextArea > div > div > textarea:focus {
        outline: none !important;
        box-shadow: 0 0 0 3px rgba(16, 185, 129, 0.2) !important;
        border: 2px solid var(--accent-green) !important;
    }
    
    /* ============= SELECTBOX COMPLETE STYLING ============= */
    /* Reset all selectbox styling first */
    .stSelectbox,
    .stSelectbox *,
    .stSelectbox > div,
    .stSelectbox > div > div,
    section[data-testid="stSidebar"] .stSelectbox,
    section[data-testid="stSidebar"] .stSelectbox *,
    section[data-testid="stSidebar"] .stSelectbox > div,
    section[data-testid="stSidebar"] .stSelectbox > div > div {
        border: none !important;
        outline: none !important;
        box-shadow: none !important;
        background: transparent !important;
    }
    
    /* Style the main selectbox container */
    .stSelectbox > div > div,
    section[data-testid="stSidebar"] .stSelectbox > div > div {
        background-color: #1A1A1A !important;
        border: 2px solid var(--accent-blue) !important;
        border-radius: 8px !important;
        padding: 8px 12px !important;
        color: var(--text-primary) !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        transition: all 0.3s ease !important;
        min-height: 40px !important;
    }
    
    /* Hover and focus states */
    .stSelectbox > div > div:hover,
    section[data-testid="stSidebar"] .stSelectbox > div > div:hover {
        border-color: var(--accent-green) !important;
        box-shadow: 0 0 0 2px rgba(16, 185, 129, 0.1) !important;
    }
    
    .stSelectbox > div > div:focus-within,
    section[data-testid="stSidebar"] .stSelectbox > div > div:focus-within {
        border-color: var(--accent-green) !important;
        box-shadow: 0 0 0 2px rgba(16, 185, 129, 0.2) !important;
    }
    
    /* Style the selected value text */
    .stSelectbox [data-baseweb="select"] > div,
    section[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] > div {
        background: transparent !important;
        border: none !important;
        color: var(--text-primary) !important;
        font-weight: 600 !important;
        padding: 0 !important;
    }
    
    /* Remove all visual artifacts */
    .stSelectbox [data-baseweb="select"],
    .stSelectbox div[role="button"],
    section[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"],
    section[data-testid="stSidebar"] .stSelectbox div[role="button"] {
        background: transparent !important;
        border: none !important;
        outline: none !important;
        box-shadow: none !important;
        color: var(--text-primary) !important;
    }
    
    /* Dropdown arrow styling */
    .stSelectbox svg,
    section[data-testid="stSidebar"] .stSelectbox svg {
        color: var(--accent-blue) !important;
        fill: var(--accent-blue) !important;
    }
    
    /* Dropdown options panel */
    .stSelectbox [role="listbox"] {
        background-color: #1A1A1A !important;
        border: 2px solid var(--accent-blue) !important;
        border-radius: 8px !important;
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.3) !important;
        margin-top: 4px !important;
        z-index: 9999 !important;
    }
    
    /* Individual dropdown options */
    .stSelectbox [role="option"] {
        background-color: transparent !important;
        color: var(--text-primary) !important;
        padding: 12px 16px !important;
        font-weight: 500 !important;
        border: none !important;
        transition: all 0.2s ease !important;
    }
    
    .stSelectbox [role="option"]:hover {
        background-color: var(--accent-blue) !important;
        color: white !important;
    }
    
    /* Remove all focus artifacts */
    .stSelectbox *:focus,
    .stSelectbox *:active,
    .stSelectbox *:focus-visible,
    section[data-testid="stSidebar"] .stSelectbox *:focus,
    section[data-testid="stSidebar"] .stSelectbox *:active,
    section[data-testid="stSidebar"] .stSelectbox *:focus-visible {
        outline: none !important;
        box-shadow: none !important;
        border: none !important;
    }
    
    /* ============= FILE UPLOADER STYLING ============= */
    .stFileUploader > div,
    .stFileUploader section,
    .stFileUploader [data-testid="stFileUploader"],
    section[data-testid="stSidebar"] .stFileUploader > div,
    section[data-testid="stSidebar"] .stFileUploader section,
    section[data-testid="stSidebar"] .stFileUploader [data-testid="stFileUploader"] {
        background-color: #1A1A1A !important;
        border: 3px dashed var(--accent-blue) !important;
        border-radius: 12px !important;
        padding: 20px !important;
        transition: all 0.3s ease !important;
        outline: none !important;
        box-shadow: var(--shadow-sm) !important;
    }
    
    .stFileUploader > div:hover,
    section[data-testid="stSidebar"] .stFileUploader > div:hover {
        border-color: var(--accent-green) !important;
        background-color: #222222 !important;
        transform: translateY(-2px) !important;
        box-shadow: var(--shadow-md) !important;
    }
    
    /* File uploader text */
    .stFileUploader span, 
    .stFileUploader p,
    .stFileUploader div,
    .stFileUploader label,
    section[data-testid="stSidebar"] .stFileUploader span,
    section[data-testid="stSidebar"] .stFileUploader p,
    section[data-testid="stSidebar"] .stFileUploader div,
    section[data-testid="stSidebar"] .stFileUploader label {
        color: var(--text-primary) !important;
        font-weight: 600 !important;
        background-color: transparent !important;
    }
    
    /* File uploader button */
    .stFileUploader button,
    section[data-testid="stSidebar"] .stFileUploader button {
        background: linear-gradient(135deg, var(--accent-blue), var(--accent-green)) !important;
        color: white !important;
        border: none !important;
        border-radius: 6px !important;
        font-weight: 600 !important;
        outline: none !important;
        box-shadow: var(--shadow-sm) !important;
        transition: all 0.3s ease !important;
        padding: 8px 16px !important;
    }
    
    .stFileUploader button:hover,
    section[data-testid="stSidebar"] .stFileUploader button:hover {
        transform: translateY(-2px) !important;
        box-shadow: var(--shadow-md) !important;
    }
    
    /* ============= BUTTON ENHANCEMENTS ============= */
    .stButton > button {
        background: linear-gradient(145deg, var(--bg-card), #374151) !important;
        color: var(--text-primary) !important;
        border: 2px solid var(--border-color) !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 10px 20px !important;
        transition: all 0.3s ease !important;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.8) !important;
        outline: none !important;
        box-shadow: var(--shadow-sm) !important;
        position: relative !important;
        overflow: hidden !important;
    }
    
    .stButton > button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.1), transparent);
        transition: left 0.5s ease;
    }
    
    .stButton > button:hover {
        background: linear-gradient(145deg, var(--accent-blue), var(--accent-green)) !important;
        border-color: var(--accent-blue) !important;
        transform: translateY(-2px) !important;
        box-shadow: var(--shadow-md) !important;
        color: white !important;
    }
    
    .stButton > button:hover::before {
        left: 100%;
    }
    
    .stButton > button:focus {
        outline: none !important;
        box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.3) !important;
    }
    
    .stButton > button:active {
        transform: translateY(0) !important;
    }
    
    /* Sidebar Buttons */
    section[data-testid="stSidebar"] .stButton > button {
        width: 100% !important;
        margin: 4px 0 !important;
    }
    
    /* Disabled buttons */
    .stButton > button:disabled {
        background: linear-gradient(145deg, #2A2A2A, #1A1A1A) !important;
        color: var(--text-muted) !important;
        border-color: var(--text-muted) !important;
        transform: none !important;
        box-shadow: none !important;
        cursor: not-allowed !important;
    }
    
    /* ============= PROGRESS BAR STYLING ============= */
    .stProgress > div > div > div {
        background: linear-gradient(90deg, var(--accent-green), var(--accent-blue)) !important;
        border-radius: 6px !important;
        box-shadow: 0 0 10px rgba(37, 99, 235, 0.5) !important;
        position: relative !important;
        overflow: hidden !important;
    }
    
    .stProgress > div > div > div::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(45deg, transparent 30%, rgba(255, 255, 255, 0.2) 50%, transparent 70%);
        animation: progress-shine 2s infinite;
    }
    
    @keyframes progress-shine {
        0% { transform: translateX(-100%); }
        100% { transform: translateX(100%); }
    }
    
    /* ============= ALERTS AND NOTIFICATIONS ============= */
    .stWarning, .stError, .stSuccess, .stInfo {
        border-radius: 8px !important;
        font-weight: 500 !important;
        border: none !important;
        box-shadow: var(--shadow-sm) !important;
        position: relative !important;
        overflow: hidden !important;
    }
    
    .stSuccess {
        background: linear-gradient(145deg, var(--success), #059669) !important;
        color: white !important;
    }
    
    .stError {
        background: linear-gradient(145deg, var(--error), #DC2626) !important;
        color: white !important;
    }
    
    .stWarning {
        background: linear-gradient(145deg, var(--warning), #D97706) !important;
        color: white !important;
    }
    
    .stInfo {
        background: linear-gradient(145deg, var(--accent-blue), #1D4ED8) !important;
        color: white !important;
    }
    
    /* ============= LABEL OVERRIDES ============= */
    .stSelectbox > label,
    .stSelectbox label,
    section[data-testid="stSidebar"] .stSelectbox > label,
    section[data-testid="stSidebar"] .stSelectbox label,
    section[data-testid="stSidebar"] label,
    .css-1d391kg label,
    .css-1d391kg .stSelectbox label,
    .css-1d391kg .stSelectbox > label {
        border: none !important;
        outline: none !important;
        box-shadow: none !important;
        background: transparent !important;
        color: var(--text-primary) !important;
        font-weight: 600 !important;
        padding: 0 !important;
        margin: 0 0 8px 0 !important;
        -webkit-appearance: none !important;
        -moz-appearance: none !important;
        appearance: none !important;
        font-size: 14px !important;
    }
    
    /* Hide Streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
    </style>
    """, unsafe_allow_html=True)

class SessionState:
    """Manage session state variables"""
    @staticmethod
    def init():
        if 'token_map' not in st.session_state:
            st.session_state.token_map = {}
        if 'scan_results' not in st.session_state:
            st.session_state.scan_results = []
        if 'scan_progress' not in st.session_state:
            st.session_state.scan_progress = 0
        if 'scan_status' not in st.session_state:
            st.session_state.scan_status = "Ready to scan"
        if 'scan_running' not in st.session_state:
            st.session_state.scan_running = False
        if 'console_messages' not in st.session_state:
            st.session_state.console_messages = ["[READY] DocXScan v3.0 initialized", 
                                                "[READY] Upload token file to begin"]
        if 'matching_files' not in st.session_state:
            st.session_state.matching_files = []
        if 'selected_folder_path' not in st.session_state:
            st.session_state.selected_folder_path = ""
        if 'folder_browser_mode' not in st.session_state:
            st.session_state.folder_browser_mode = "select"  # "select", "input", "browse"
        if 'path_history' not in st.session_state:
            st.session_state.path_history = []
        if 'current_path_input' not in st.session_state:
            st.session_state.current_path_input = ""

class DocumentScanner:
    """Core document scanning functionality"""
    
    @staticmethod
    def extract_full_text_lines(doc):
        """Extract text from document"""
        lines = []
        try:
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text)
        except Exception as e:
            lines.append(f"Error extracting text: {str(e)}")
        return lines
    
    @staticmethod
    def scan_documents(folder_path, patterns, file_filter, progress_placeholder, console_placeholder):
        """Main document scanning logic"""
        try:
            matching_files = []
            metadata = []
            
            # Log start
            log_message("🔍 Starting document scan...", console_placeholder)
            log_message(f"📂 Folder: {os.path.basename(folder_path)}", console_placeholder)
            log_message(f"🎯 Patterns: {', '.join(patterns[:3])}{'...' if len(patterns) > 3 else ''}", console_placeholder)
            
            # Collect all files
            all_files = []
            for root_dir, _, files in os.walk(folder_path):
                for file in files:
                    if file_filter(file) and not file.startswith('~'):
                        full_path = os.path.join(root_dir, file)
                        if os.path.exists(full_path):
                            all_files.append(full_path)
            
            if not all_files:
                log_message("❌ No files found to scan", console_placeholder)
                return [], []
            
            log_message(f"📄 Found {len(all_files)} files to process", console_placeholder)
            
            # Process files
            for i, full_path in enumerate(all_files):
                try:
                    progress = int((i / len(all_files)) * 100)
                    filename = os.path.basename(full_path)
                    
                    # Update progress
                    st.session_state.scan_progress = progress
                    st.session_state.scan_status = f"Processing {filename[:20]}..."
                    progress_placeholder.progress(progress / 100)
                    
                    # Process document
                    doc = Document(full_path)
                    full_text = '\n'.join(DocumentScanner.extract_full_text_lines(doc))
                    matched = []
                    matched_lines = []
                    
                    # Check for patterns
                    for token in patterns:
                        if token in full_text:
                            matched.append(token)
                            # Find matching lines
                            for line in full_text.split('\n'):
                                if token in line and line.strip():
                                    matched_lines.append(line.strip()[:100])  # Limit line length
                    
                    if matched:
                        matching_files.append(full_path)
                        
                        # Get file info
                        try:
                            info = os.stat(full_path)
                            file_size = info.st_size
                            creation_date = datetime.fromtimestamp(info.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
                            modified_date = datetime.fromtimestamp(info.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                        except Exception:
                            file_size = 0
                            creation_date = "Unknown"
                            modified_date = "Unknown"
                        
                        token_count = sum(full_text.count(token) for token in matched)
                        
                        metadata.append({
                            'File Name': filename,
                            'File Path': full_path,
                            'Size (bytes)': file_size,
                            'Creation Date': creation_date,
                            'Modified Date': modified_date,
                            'Matched Pattern(s)': ', '.join(matched),
                            'Matched Line(s)': ' | '.join(matched_lines[:3]),  # Limit to 3 lines
                            'Token Match Count': token_count
                        })
                        
                        log_message(f"✅ Match found: {filename}", console_placeholder)
                
                except Exception as e:
                    log_message(f"❌ Error processing {filename}: {str(e)}", console_placeholder)
                
                # Small delay to allow UI updates
                time.sleep(0.01)
            
            # Complete
            st.session_state.scan_progress = 100
            st.session_state.scan_status = "Scan completed!"
            progress_placeholder.progress(1.0)
            
            if matching_files:
                log_message(f"🎉 Scan complete! Found {len(matching_files)} matching files", console_placeholder)
            else:
                log_message("ℹ️ No matching files found", console_placeholder)
            
            return matching_files, metadata
            
        except Exception as e:
            log_message(f"❌ Scan failed: {str(e)}", console_placeholder)
            return [], []

def log_message(message, console_placeholder=None):
    """Add message to console log"""
    timestamp = datetime.now().strftime("%H:%M:%S")
    formatted_msg = f"[{timestamp}] {message}"
    st.session_state.console_messages.append(formatted_msg)
    
    # Keep only last 30 messages
    if len(st.session_state.console_messages) > 30:
        st.session_state.console_messages = st.session_state.console_messages[-30:]
    
    # Update console display if placeholder provided
    if console_placeholder:
        console_text = '\n'.join(st.session_state.console_messages)
        console_placeholder.markdown(
            f'<div class="console-area">{console_text}</div>',
            unsafe_allow_html=True
        )

def clear_console():
    """Clear console messages"""
    st.session_state.console_messages = ["[READY] Console cleared"]

def create_template():
    """Create token template JSON"""
    template = {
        "<<FileService.": "Fileservice",
        "</ff>": "Page Break",
        "</pp>": "Hard Return",
        "<backspace>": "Backspace",
        "<<STNDRDTH": "STNDRD Add \"TH\"",
        "<c>": "Center",
        "<u>": "Underline",
        "<i>": "Italic",
        "<pcase>": "pcase",
        "<lcase>": "lcase",
        "<ucase>": "ucase",
        "<bold>": "Bold",
        "<nobullet>": "No Bullet",
        "<fontsize": "Font Size",
        "<s1>": "s1",
        "<s2>": "s2",
        "[[MCOMPUTEINTO(<<": "MCOMPUTE INTO",
        "[[SCOMPUTEINTO(": "SCOMPUTE INTO",
        "[[ABORTIIF": "Abortif",
        "PROMTINTO(": "PROMTINTO",
        "PROMTINTOIIF(": "PROMTINTOIIF",
        "PROMTINTOLIST(": "PROMTINTOLIST",
        "PROMTINTOIIFLIST(": "PROMTINTOIIFLIST",
        "PROMTFORM(": "PROMTFORM",
        "<<Checklist.": "CHECKLIST",
        "TABLE(": "TABLE",
        "<<jfig": "JFIG",
        "jfig": "JFIG_General",
        "{ATTY": "ESIGN",
        "<<Special.": "SPECIAL",
        "+91|<<Special.ToDay": "+91 special day",
        "-91|<<Special.ToDay": "-91 special day",
        "+2|<<Special.ToDay": "+2 special day",
        "-2|<<Special.ToDay": "-2 special day",
        "<<Tracker.MortDate>>~MMMM dd": "MMMM dd,yyyy",
        "<<Tracker.MortDate>>~MM-dd-yyyy": "MM-dd-yyyy",
        "<<Tracker.MortDate>>~ddd": "ddd,MMM dd-yyyy",
        "<<Tracker.OriginalPrincipal>>~##": "##,###,###.00",
        "CU$TOMMMMMMPLACEHOLDER": "Custom(Enter Below)"
    }
    
    return json.dumps(template, indent=2)

def create_zip_download(matching_files, metadata, zip_name="matched_files"):
    """Create ZIP file for download"""
    try:
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Create Excel metadata file
            excel_buffer = BytesIO()
            df = pd.DataFrame(metadata)
            df.to_excel(excel_buffer, index=False, engine='openpyxl')
            excel_buffer.seek(0)
            
            # Add Excel file to ZIP
            zipf.writestr('scan_results.xlsx', excel_buffer.getvalue())
            
            # Add matched files
            for file_path in matching_files:
                if os.path.exists(file_path):
                    arcname = os.path.join('matched_files', os.path.basename(file_path))
                    zipf.write(file_path, arcname)
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error creating ZIP: {str(e)}")
        return None

def format_file_size(size_bytes):
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    return f"{size_bytes:.1f} {size_names[i]}"

def get_drives_windows():
    """Get available drives on Windows"""
    drives = []
    for letter in 'ABCDEFGHIJKLMNOPQRSTUVWXYZ':
        drive = f"{letter}:\\"
        if os.path.exists(drive):
            drives.append(drive)
    return drives

def get_recent_folders():
    """Get recently accessed folders (simple heuristic)"""
    recent = []
    user_home = os.path.expanduser("~")
    
    # Common recent folder locations
    potential_folders = [
        os.path.join(user_home, "Documents"),
        os.path.join(user_home, "Desktop"),
        os.path.join(user_home, "Downloads"),
        os.path.join(user_home, "OneDrive"),
        os.path.join(user_home, "Google Drive"),
        os.path.join(user_home, "Dropbox"),
    ]
    
    for folder in potential_folders:
        if os.path.exists(folder):
            # Check if it contains any .docx files
            try:
                docx_count = len(glob.glob(os.path.join(folder, "**/*.docx"), recursive=True))
                if docx_count > 0:
                    recent.append((folder, docx_count))
            except:
                pass
    
    return recent

def smart_path_suggestions(current_input):
    """Provide intelligent path suggestions based on current input"""
    if not current_input:
        return []
    
    try:
        # If it's a partial path, try to complete it
        if os.path.sep in current_input:
            parent = os.path.dirname(current_input)
            if os.path.exists(parent):
                basename = os.path.basename(current_input).lower()
                items = []
                
                for item in os.listdir(parent):
                    full_path = os.path.join(parent, item)
                    if (os.path.isdir(full_path) and 
                        item.lower().startswith(basename) and
                        not item.startswith('.')):
                        # Count docx files in this directory
                        try:
                            docx_count = len(glob.glob(os.path.join(full_path, "**/*.docx"), recursive=True))
                            items.append((full_path, item, docx_count))
                        except:
                            items.append((full_path, item, 0))
                
                # Sort by docx count (descending) and name
                items.sort(key=lambda x: (-x[2], x[1].lower()))
                return items[:8]  # Limit to 8 suggestions
        
        # If it's just a drive letter or root, show top-level directories
        elif current_input.endswith(":\\") or current_input == "/":
            items = []
            try:
                for item in os.listdir(current_input):
                    full_path = os.path.join(current_input, item)
                    if os.path.isdir(full_path) and not item.startswith('.'):
                        try:
                            docx_count = len(glob.glob(os.path.join(full_path, "**/*.docx"), recursive=True))
                            if docx_count > 0:  # Only show folders with documents
                                items.append((full_path, item, docx_count))
                        except:
                            pass
                
                items.sort(key=lambda x: (-x[2], x[1].lower()))
                return items[:8]
            except PermissionError:
                return []
    
    except Exception:
        pass
    
    return []

def render_super_enhanced_folder_browser():
    """Super enhanced folder browser with intelligent features"""
    
    if st.session_state.folder_browser_mode == "select":
        # Main selection interface
        st.markdown("""
        <div class="folder-browser-enhanced">
            <div class="folder-browser-icon">📁</div>
            <div class="folder-browser-title">Smart Folder Selection</div>
            <div class="folder-browser-subtitle">Multiple ways to find your document folder</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Create tabs for different selection methods
        tab1, tab2, tab3, tab4 = st.tabs(["🚀 Quick Access", "📍 Recent & Smart", "💾 Drives", "📝 Manual Entry"])
        
        with tab1:
            st.markdown("#### 🏠 Common Locations")
            
            # Get system info
            user_home = os.path.expanduser("~")
            system = platform.system()
            
            # Quick access locations
            quick_locations = [
                ("📄 Documents", os.path.join(user_home, "Documents")),
                ("💻 Desktop", os.path.join(user_home, "Desktop")),
                ("📁 Downloads", os.path.join(user_home, "Downloads")),
                ("👤 Home Folder", user_home),
            ]
            
            # Add cloud storage locations if they exist
            cloud_locations = [
                ("☁️ OneDrive", os.path.join(user_home, "OneDrive")),
                ("📂 Google Drive", os.path.join(user_home, "Google Drive")),
                ("📦 Dropbox", os.path.join(user_home, "Dropbox")),
                ("☁️ iCloud", os.path.join(user_home, "iCloud Drive")),
            ]
            
            for name, path in cloud_locations:
                if os.path.exists(path):
                    quick_locations.append((name, path))
            
            # Display quick access buttons
            cols = st.columns(2)
            for i, (name, path) in enumerate(quick_locations):
                col = cols[i % 2]
                with col:
                    if os.path.exists(path):
                        # Count documents in this location
                        try:
                            docx_count = len(glob.glob(os.path.join(path, "**/*.docx"), recursive=True))
                            button_text = f"{name}"
                            if docx_count > 0:
                                button_text += f" ({docx_count} docs)"
                        except:
                            button_text = name
                            docx_count = 0
                        
                        if st.button(button_text, key=f"quick_{i}", use_container_width=True):
                            st.session_state.selected_folder_path = path
                            if path not in st.session_state.path_history:
                                st.session_state.path_history.insert(0, path)
                            log_message(f"📂 Quick select: {path}")
                            st.rerun()
                    else:
                        st.button(f"{name} (Not Found)", disabled=True, use_container_width=True, key=f"quick_disabled_{i}")
        
        with tab2:
            st.markdown("#### 🕒 Smart Suggestions")
            
            # Get folders with documents
            recent_folders = get_recent_folders()
            
            if recent_folders:
                st.markdown("**Folders containing documents:**")
                for i, (folder, count) in enumerate(recent_folders[:6]):
                    folder_name = os.path.basename(folder) or folder
                    if st.button(f"📁 {folder_name} ({count} docs)", key=f"recent_{i}", use_container_width=True):
                        st.session_state.selected_folder_path = folder
                        if folder not in st.session_state.path_history:
                            st.session_state.path_history.insert(0, folder)
                        log_message(f"📂 Selected from suggestions: {folder}")
                        st.rerun()
            else:
                st.info("No folders with .docx files found in common locations")
            
            # Show path history if available
            if st.session_state.path_history:
                st.markdown("**📚 Previously Used:**")
                for i, path in enumerate(st.session_state.path_history[:5]):
                    if os.path.exists(path):
                        path_name = os.path.basename(path) or path
                        if st.button(f"🕒 {path_name}", key=f"history_{i}", use_container_width=True):
                            st.session_state.selected_folder_path = path
                            log_message(f"📂 Selected from history: {path}")
                            st.rerun()
        
        with tab3:
            st.markdown("#### 💾 Browse Drives")
            
            if platform.system() == "Windows":
                drives = get_drives_windows()
                st.markdown("**Available Drives:**")
                
                drive_cols = st.columns(min(4, len(drives)))
                for i, drive in enumerate(drives):
                    col = drive_cols[i % len(drive_cols)]
                    with col:
                        try:
                            drive_label = f"💾 {drive}"
                            
                            # Count documents on this drive
                            docx_count = len(glob.glob(os.path.join(drive, "**/*.docx"), recursive=True))
                            if docx_count > 0:
                                drive_label += f" ({docx_count})"
                            
                        except:
                            drive_label = f"💾 {drive}"
                        
                        if st.button(drive_label, key=f"drive_{i}", use_container_width=True):
                            st.session_state.current_path_input = drive
                            st.session_state.folder_browser_mode = "browse"
                            st.rerun()
            else:
                # Unix-like systems
                common_roots = ["/", "/home", "/Users", "/mnt", "/media"]
                for i, root in enumerate(common_roots):
                    if os.path.exists(root):
                        if st.button(f"📁 {root}", key=f"root_{root.replace('/', '_')}", use_container_width=True):
                            st.session_state.current_path_input = root
                            st.session_state.folder_browser_mode = "browse"
                            st.rerun()
        
        with tab4:
            st.markdown("#### 📝 Enter Path Manually")
            if st.button("✏️ Open Manual Entry", use_container_width=True, key="open_manual"):
                st.session_state.folder_browser_mode = "input"
                st.rerun()
        
        # Help section
        with st.expander("❓ Need Help Finding Your Folder?", expanded=False):
            st.markdown(f"""
            **🖥️ System Detected:** {platform.system()} {platform.release()}
            
            **💡 Tips:**
            - Use **Quick Access** for common locations
            - Check **Smart Suggestions** for folders containing documents
            - **Drives** tab shows all available storage devices
            - **Manual Entry** for custom paths
            
            **🔍 Looking for specific file types?**
            - We'll scan for `.docx` and `.dcp.docx` files
            - Subfolders are included automatically
            - Hidden files (starting with . or ~) are ignored
            """)
    
    elif st.session_state.folder_browser_mode == "browse":
        render_folder_browser_interface()
    
    elif st.session_state.folder_browser_mode == "input":
        render_manual_path_input()

def render_folder_browser_interface():
    """Interactive folder browsing interface"""
    current_path = st.session_state.current_path_input
    
    st.markdown("### 🗂️ Browse Folders")
    
    # Breadcrumb navigation
    if current_path:
        path_parts = Path(current_path).parts
        
        cols = st.columns(len(path_parts) + 1)
        
        # Root/drive
        with cols[0]:
            if st.button("🏠", key="nav_root", help="Go to root"):
                if platform.system() == "Windows":
                    st.session_state.current_path_input = path_parts[0] + "\\"
                else:
                    st.session_state.current_path_input = "/"
                st.rerun()
        
        # Path parts
        for i, part in enumerate(path_parts):
            if i + 1 < len(cols):
                with cols[i + 1]:
                    partial_path = str(Path(*path_parts[:i+1]))
                    if platform.system() == "Windows" and i == 0:
                        partial_path += "\\"
                    
                    if st.button(part[:10] + ("..." if len(part) > 10 else ""), 
                               key=f"nav_{i}", 
                               help=f"Go to {partial_path}"):
                        st.session_state.current_path_input = partial_path
                        st.rerun()
    
    # Current directory contents
    try:
        if os.path.exists(current_path) and os.path.isdir(current_path):
            items = []
            
            # Add parent directory option
            parent = os.path.dirname(current_path)
            if parent != current_path:
                items.append(("📁 ..", parent, True, 0))
            
            # List directories and count documents
            for item in sorted(os.listdir(current_path)):
                if not item.startswith('.'):
                    full_path = os.path.join(current_path, item)
                    if os.path.isdir(full_path):
                        try:
                            docx_count = len(glob.glob(os.path.join(full_path, "*.docx")))
                            items.append((f"📁 {item}", full_path, True, docx_count))
                        except:
                            items.append((f"📁 {item}", full_path, True, 0))
            
            # Display items
            if items:
                st.markdown(f"**Contents of:** `{current_path}`")
                
                # Show folders in columns
                folder_cols = st.columns(2)
                for i, (display_name, full_path, is_dir, doc_count) in enumerate(items):
                    col = folder_cols[i % 2]
                    with col:
                        button_text = display_name
                        if doc_count > 0:
                            button_text += f" ({doc_count} docs)"
                        
                        if st.button(button_text, key=f"browse_item_{i}", use_container_width=True):
                            st.session_state.current_path_input = full_path
                            st.rerun()
            else:
                st.info("📂 Empty directory")
            
            # Action buttons
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("✅ Select This Folder", use_container_width=True, key="select_browsed_folder"):
                    st.session_state.selected_folder_path = current_path
                    st.session_state.folder_browser_mode = "select"
                    if current_path not in st.session_state.path_history:
                        st.session_state.path_history.insert(0, current_path)
                    log_message(f"📂 Folder selected via browse: {current_path}")
                    st.rerun()
            
            with col2:
                if st.button("🔄 Refresh", use_container_width=True, key="refresh_browse"):
                    st.rerun()
            
            with col3:
                if st.button("🔙 Back to Quick Access", use_container_width=True, key="back_to_quick"):
                    st.session_state.folder_browser_mode = "select"
                    st.rerun()
        
        else:
            st.error("❌ Invalid path or permission denied")
            if st.button("🔙 Back", use_container_width=True, key="back_from_error"):
                st.session_state.folder_browser_mode = "select"
                st.rerun()
    
    except Exception as e:
        st.error(f"❌ Error browsing folder: {str(e)}")
        if st.button("🔙 Back", use_container_width=True, key="back_from_exception"):
            st.session_state.folder_browser_mode = "select"
            st.rerun()

def render_manual_path_input():
    """Enhanced manual path input with smart suggestions"""
    st.markdown("### ✏️ Manual Path Entry")
    
    # Path input with live suggestions
    current_input = st.text_input(
        "📂 Folder Path",
        value=st.session_state.current_path_input,
        placeholder="Start typing: C:\\Users\\... or /home/user/...",
        help="Type your folder path and get smart suggestions",
        key="manual_path_input_smart"
    )
    
    # Update session state
    if current_input != st.session_state.current_path_input:
        st.session_state.current_path_input = current_input
    
    # Smart path suggestions
    if current_input:
        suggestions = smart_path_suggestions(current_input)
        
        if suggestions:
            st.markdown("**💡 Smart Suggestions:**")
            
            suggestion_cols = st.columns(min(2, len(suggestions)))
            for i, (full_path, item_name, doc_count) in enumerate(suggestions):
                col = suggestion_cols[i % len(suggestion_cols)]
                with col:
                    display_text = f"📁 {item_name}"
                    if doc_count > 0:
                        display_text += f" ({doc_count} docs)"
                    
                    if st.button(display_text, key=f"suggest_manual_{i}", use_container_width=True):
                        st.session_state.current_path_input = full_path
                        st.rerun()
    
    # Path validation
    path_valid = False
    if current_input:
        if os.path.exists(current_input) and os.path.isdir(current_input):
            try:
                # Count documents
                docx_files = glob.glob(os.path.join(current_input, "**/*.docx"), recursive=True)
                dcp_files = glob.glob(os.path.join(current_input, "**/*.dcp.docx"), recursive=True)
                total_docs = len(docx_files) + len(dcp_files)
                
                st.success(f"✅ Valid folder with {total_docs} document files")
                path_valid = True
                
                # Show some file examples
                if total_docs > 0:
                    with st.expander(f"📄 Preview ({min(5, total_docs)} of {total_docs} files)", expanded=False):
                        all_files = docx_files + dcp_files
                        for file_path in all_files[:5]:
                            file_name = os.path.basename(file_path)
                            st.caption(f"📄 {file_name}")
                
            except Exception as e:
                st.warning(f"⚠️ Valid path but couldn't scan: {str(e)}")
                path_valid = True
        elif os.path.exists(current_input):
            st.error("❌ Path exists but is not a directory")
        else:
            st.error("❌ Path does not exist")
    
    # Action buttons
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("✅ Use This Path", disabled=not path_valid, use_container_width=True, key="confirm_manual_path"):
            if path_valid:
                st.session_state.selected_folder_path = current_input
                st.session_state.folder_browser_mode = "select"
                if current_input not in st.session_state.path_history:
                    st.session_state.path_history.insert(0, current_input)
                log_message(f"📂 Manual path confirmed: {current_input}")
                st.rerun()
    
    with col2:
        if st.button("🗂️ Browse Mode", use_container_width=True, key="switch_to_browse"):
            if current_input and os.path.exists(current_input):
                st.session_state.current_path_input = current_input
            st.session_state.folder_browser_mode = "browse"
            st.rerun()
    
    with col3:
        if st.button("🔙 Back", use_container_width=True, key="back_from_manual"):
            st.session_state.folder_browser_mode = "select"
            st.rerun()
    
    # OS-specific help
    with st.expander("📖 Path Format Help", expanded=False):
        system = platform.system()
        
        if system == "Windows":
            st.markdown("""
            **Windows Path Examples:**
            - `C:\\Users\\YourName\\Documents`
            - `D:\\Projects\\Legal Documents`
            - `\\\\NetworkDrive\\SharedFolder`
            
            **Tips:**
            - Use double backslashes `\\\\` or forward slashes `/`
            - Drive letters are case-insensitive
            - Network paths start with `\\\\`
            """)
        else:
            st.markdown("""
            **Unix/Linux/macOS Path Examples:**
            - `/home/username/Documents`
            - `/Users/username/Projects`
            - `/mnt/external/WorkFiles`
            
            **Tips:**
            - Paths are case-sensitive
            - Use forward slashes `/`
            - Home directory: `~` or `/home/username`
            """)

def main():
    """Main application"""
    load_css()
    SessionState.init()
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1 class="main-title">🔍 DocXScan v3.0</h1>
        <p class="main-subtitle">Professional document scanner with intelligent token detection</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar Configuration
    with st.sidebar:
        st.markdown("### ⚙️ Configuration")
        
        # Token file upload with enhanced styling
        st.markdown("#### 📂 Token Management")
        uploaded_token_file = st.file_uploader(
            "Upload Token JSON File",
            type=['json'],
            help="Upload your token mappings file",
            label_visibility="visible"
        )
        
        if uploaded_token_file is not None:
            try:
                token_data = json.load(uploaded_token_file)
                st.session_state.token_map = token_data
                st.success(f"✅ Loaded {len(token_data)} tokens")
                log_message(f"✅ Loaded {len(token_data)} tokens from {uploaded_token_file.name}")
            except Exception as e:
                st.error(f"❌ Error loading token file: {str(e)}")
                log_message(f"❌ Token load error: {str(e)}")
        
        # Token selection
        st.markdown("#### 🎯 Token Selection")
        if st.session_state.token_map:
            token_options = ["-- Select Token --"] + sorted(list(st.session_state.token_map.values()))
            selected_token = st.selectbox(
                "Choose Token",
                token_options,
                help="Choose from loaded tokens",
                key="token_selector"
            )
        else:
            selected_token = "-- Select Token --"
            st.warning("⚠️ Please upload a token file first")
        
        # Custom tokens
        st.markdown("#### ✏️ Custom Tokens")
        custom_tokens = st.text_area(
            "Additional Tokens",
            help="Enter additional tokens separated by commas",
            placeholder="<<token1>>, [[token2]], {token3}",
            height=80,
            key="custom_tokens_input"
        )
        
        if custom_tokens:
            token_count = len([t.strip() for t in custom_tokens.split(",") if t.strip()])
            st.caption(f"📝 {token_count} custom tokens added")
        
        # File type filter
        st.markdown("#### 📁 File Filter")
        file_type = st.selectbox(
            "File Types",
            ["Both (.docx and .dcp.docx)", "Only .dcp.docx", "Only .docx (excluding .dcp.docx)"],
            help="Select which file types to scan",
            key="file_type_selector"
        )
        
        # ZIP output name
        st.markdown("#### 📦 Output Settings")
        zip_name = st.text_input(
            "ZIP File Name",
            value="matched_files",
            help="Name for the output ZIP file",
            key="zip_name_input"
        )
        
        st.markdown("---")
        
        # Action buttons
        st.markdown("#### 🛠️ Actions")
        
        # Template creation
        if st.button("📄 Create Template", use_container_width=True, key="create_template_btn"):
            template_json = create_template()
            st.download_button(
                label="💾 Download Template",
                data=template_json,
                file_name="token_template.json",
                mime="application/json",
                use_container_width=True,
                key="download_template_btn"
            )
            log_message("📄 Token template created")
        
        # Clear console
        if st.button("🧹 Clear Console", use_container_width=True, key="clear_console_btn"):
            clear_console()
    
    # Main content area
    col1, col2 = st.columns([2, 1], gap="medium")
    
    with col1:
        # Folder selection - Super Enhanced browser
        st.markdown("""
        <div class="modern-card">
            <div class="card-title">📁 Advanced Folder Selection <div class="status-indicator"></div></div>
        </div>
        """, unsafe_allow_html=True)
        
        # Check if folder is already selected
        if not st.session_state.selected_folder_path:
            render_super_enhanced_folder_browser()  # NEW ENHANCED FUNCTION
            folder_valid = False
            folder_path = ""
        else:
            # Show selected folder with file count
            file_filter_map = {
                "Only .dcp.docx": lambda f: f.endswith('.dcp.docx'),
                "Only .docx (excluding .dcp.docx)": lambda f: f.endswith('.docx') and not f.endswith('.dcp.docx'),
                "Both (.docx and .dcp.docx)": lambda f: f.endswith('.docx')
            }
            file_filter = file_filter_map.get(file_type, lambda f: f.endswith('.docx'))
            
            # Count files using the enhanced method
            file_count = 0
            try:
                docx_files = glob.glob(os.path.join(st.session_state.selected_folder_path, "**/*.docx"), recursive=True)
                file_count = len([f for f in docx_files if file_filter(os.path.basename(f))])
            except:
                file_count = 0
            
            # Display selected folder
            st.markdown(f"""
            <div class="selected-folder-display">
                <div class="selected-folder-text">✅ Selected Folder</div>
                <div class="selected-folder-path">{st.session_state.selected_folder_path}</div>
                <div style="margin-top: 10px; font-size: 0.9rem;">
                    📄 Found {file_count} document files to scan
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Change folder button
            if st.button("🔄 Change Folder", use_container_width=True, key="change_folder_btn"):
                st.session_state.selected_folder_path = ""
                st.session_state.folder_browser_mode = "select"
                st.rerun()
            
            folder_valid = True
            folder_path = st.session_state.selected_folder_path
        
        # Scan controls
        st.markdown("""
        <div class="modern-card">
            <div class="card-title">⚡ Scan Controls <div class="status-indicator"></div></div>
        </div>
        """, unsafe_allow_html=True)
        
        # Progress display
        progress_placeholder = st.empty()
        status_placeholder = st.empty()
        
        if st.session_state.scan_progress > 0:
            progress_placeholder.progress(st.session_state.scan_progress / 100)
            status_placeholder.info(f"Status: {st.session_state.scan_status}")
        
        # Check if scan can be started
        can_scan = (
            st.session_state.selected_folder_path and 
            (st.session_state.token_map or custom_tokens) and 
            not st.session_state.scan_running
        )
        
        col_btn1, col_btn2, col_btn3 = st.columns([2, 1, 1])
        
        with col_btn1:
            if st.button("🚀 Start Scan", disabled=not can_scan, use_container_width=True, key="start_scan_btn"):
                if can_scan:
                    # Prepare patterns
                    patterns = []
                    
                    # Add selected token
                    if selected_token != "-- Select Token --":
                        matched_tokens = [k for k, v in st.session_state.token_map.items() if v == selected_token]
                        patterns.extend(matched_tokens)
                    
                    # Add custom tokens
                    if custom_tokens:
                        custom_list = [t.strip() for t in custom_tokens.split(",") if t.strip()]
                        patterns.extend(custom_list)
                    
                    if patterns:
                        # Define file filter
                        file_filter_map = {
                            "Only .dcp.docx": lambda f: f.endswith('.dcp.docx'),
                            "Only .docx (excluding .dcp.docx)": lambda f: f.endswith('.docx') and not f.endswith('.dcp.docx'),
                            "Both (.docx and .dcp.docx)": lambda f: f.endswith('.docx')
                        }
                        file_filter = file_filter_map.get(file_type, lambda f: f.endswith('.docx'))
                        
                        # Start scan
                        st.session_state.scan_running = True
                        
                        # Create placeholders for real-time updates
                        console_placeholder = st.empty()
                        
                        # Execute scan
                        matching_files, metadata = DocumentScanner.scan_documents(
                            st.session_state.selected_folder_path,
                            patterns,
                            file_filter,
                            progress_placeholder,
                            console_placeholder
                        )
                        
                        # Store results
                        st.session_state.scan_results = metadata
                        st.session_state.matching_files = matching_files
                        st.session_state.scan_running = False
                        
                        if matching_files:
                            st.success(f"🎉 Scan completed! Found {len(matching_files)} matching files")
                            st.balloons()
                        else:
                            st.info("ℹ️ Scan completed but no matching files were found")
        
        with col_btn2:
            if st.button("📊 Results", use_container_width=True, key="results_btn"):
                if st.session_state.scan_results:
                    st.info(f"📊 {len(st.session_state.scan_results)} files in results")
                else:
                    st.warning("No results available")
        
        with col_btn3:
            if st.button("🔄 Reset", use_container_width=True, key="reset_btn"):
                st.session_state.scan_results = []
                st.session_state.matching_files = []
                st.session_state.scan_progress = 0
                st.session_state.scan_status = "Ready to scan"
                st.session_state.selected_folder_path = ""
                st.session_state.folder_browser_mode = "select"
                clear_console()
                st.rerun()
        
        # Results section
        if st.session_state.scan_results:
            st.markdown("""
            <div class="modern-card">
                <div class="card-title">📊 Scan Results <div class="status-indicator"></div></div>
            </div>
            """, unsafe_allow_html=True)
            
            # Metrics
            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            
            with col_m1:
                st.metric("📄 Files Found", len(st.session_state.scan_results))
            
            with col_m2:
                total_matches = sum(result.get('Token Match Count', 0) for result in st.session_state.scan_results)
                st.metric("🎯 Total Matches", total_matches)
            
            with col_m3:
                total_size = sum(result.get('Size (bytes)', 0) for result in st.session_state.scan_results)
                st.metric("💾 Total Size", format_file_size(total_size))
            
            with col_m4:
                unique_patterns = len(set(result.get('Matched Pattern(s)', '') for result in st.session_state.scan_results if result.get('Matched Pattern(s)')))
                st.metric("🔍 Unique Patterns", unique_patterns)
            
            # Download buttons
            col_dl1, col_dl2 = st.columns(2)
            
            with col_dl1:
                if st.session_state.matching_files:
                    zip_data = create_zip_download(st.session_state.matching_files, st.session_state.scan_results, zip_name)
                    if zip_data:
                        st.download_button(
                            label="📦 Download ZIP Package",
                            data=zip_data,
                            file_name=f"{zip_name}.zip",
                            mime="application/zip",
                            use_container_width=True,
                            key="download_zip_btn"
                        )
            
            with col_dl2:
                # Excel export
                excel_buffer = BytesIO()
                df = pd.DataFrame(st.session_state.scan_results)
                df.to_excel(excel_buffer, index=False, engine='openpyxl')
                excel_buffer.seek(0)
                
                st.download_button(
                    label="📊 Download Excel Report",
                    data=excel_buffer.getvalue(),
                    file_name=f"{zip_name}_report.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="download_excel_btn"
                )
            
            # Detailed results
            with st.expander("📋 Detailed Results", expanded=False):
                # Display columns selection
                all_columns = list(st.session_state.scan_results[0].keys())
                display_columns = ['File Name', 'Matched Pattern(s)', 'Token Match Count', 'Size (bytes)', 'Modified Date']
                available_columns = [col for col in display_columns if col in all_columns]
                
                if available_columns:
                    results_df = pd.DataFrame(st.session_state.scan_results)
                    st.dataframe(
                        results_df[available_columns],
                        use_container_width=True,
                        height=400
                    )
                else:
                    results_df = pd.DataFrame(st.session_state.scan_results)
                    st.dataframe(results_df, use_container_width=True, height=400)
    
    with col2:
        # Console section
        st.markdown("""
        <div class="modern-card">
            <div class="card-title">📋 Console Output <div class="status-indicator"></div></div>
        </div>
        """, unsafe_allow_html=True)
        
        # Console display
        console_text = '\n'.join(st.session_state.console_messages[-15:])  # Show last 15 messages
        st.markdown(
            f'<div class="console-area">{console_text}</div>',
            unsafe_allow_html=True
        )
        
        # System information
        st.markdown("### 💻 System Status")
        
        status_info = {
            "🐍 Python Version": f"{os.sys.version_info.major}.{os.sys.version_info.minor}.{os.sys.version_info.micro}",
            "⏰ Current Time": datetime.now().strftime('%H:%M:%S'),
            "📊 Console Lines": len(st.session_state.console_messages),
            "🔧 Tokens Loaded": len(st.session_state.token_map),
            "📄 Results Cached": len(st.session_state.scan_results)
        }
        
        for label, value in status_info.items():
            col_info1, col_info2 = st.columns([2, 1])
            with col_info1:
                st.caption(label)
            with col_info2:
                st.caption(f"**{value}**")
    
    # Footer
    st.markdown("---")
    st.markdown(
        '<div style="text-align: center; color: var(--text-muted); font-size: 0.875rem; padding: 1rem;">'
        '© 2025 Hrishik Kunduru • DocXScan v3.0 Professional • All Rights Reserved'
        '</div>',
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()